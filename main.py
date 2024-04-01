import threading
import requests
from bs4 import BeautifulSoup
from loguru import logger
import re
import json
import os
from io import BytesIO
from datetime import datetime
from threading import Thread

from docx import *
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt,Cm

# global variables
USER_AGENT = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Safari/537.36'
TASK_FILE = 'task.txt'
THREADS = 1
OUTPUT_PATH = 'output'

# debug option
DEBUG_DUMP_HTML = True
DEBUG_DUMP_JSON = True

def save_to_file(file, content):
    with open(file, 'w', encoding='utf-8') as f:
        f.write(content)

def load_from_file(file):
    with open(file, 'r', encoding='utf-8') as f:
        return f.read()

def download_image(url):
    response = requests.get(url, headers={'User-Agent': USER_AGENT})
    if response.status_code != 200:
        logger.error(f'状态码：{response.status_code}')
        logger.error(response.text)
        raise RuntimeError(f'图片下载失败')
    return BytesIO(response.content)

def get_html(url):
    response = requests.get(url, headers={'User-Agent': USER_AGENT})
    if response.status_code != 200:
        logger.error(f'状态码：{response.status_code}')
        logger.error(response.text)
        raise RuntimeError('网页访问失败')
    return response.text

def get_article(html):
    soup = BeautifulSoup(html, 'html.parser')
    script = soup.find(name='script').text

    res = re.search(r'var ARTICLE_DETAIL = (.*);\n    var', script, re.DOTALL)
    j = json.loads(res.group(1))
    return j

def get_title(article):
    title = article['article']['title']
    return title

def get_create_time(article):
    create_time = article['article']['create_time']
    return datetime.fromtimestamp(float(create_time))

def get_last_modify_time(article):
    last_modify_time = article['article']['last_modify_time']
    return datetime.fromtimestamp(float(last_modify_time))

def get_cover_url(article):
    cover_img_url = article['article']['cover_img_url']
    return cover_img_url

def get_author_name(article):
    nickname = article['author']['nickname']
    return nickname

def get_author_avatar_url(article):
    head_img_url = article['author']['head_img_url']
    return head_img_url

def get_author_location(article):
    country = article['author']['country']
    province = article['author']['province']
    city = article['author']['city']
    return f'{country}{province}省{city}市'

def get_author_ip_when_write(article):
    ip = article['article']['ext']['ip']
    return ip

def get_author_ip_when_login(article):
    ip = article['author']['ip_address']
    return ip

def get_author_phone_number(article):
    phone_num = article['author']['phone_num']
    return phone_num

def get_author_wechat_id(article):
    wechat_id = article['author']['wechat_id']
    return wechat_id

def get_music_name(article):
    music_desc = article['article']['music_desc']
    return music_desc

def get_music_url(article):
    music_url = article['article']['music_url']
    return music_url

def get_content(article):
    return article['article']['content']['content']

def convert_to_plaintext_array(text):
    soup = BeautifulSoup(text, 'html.parser')
    return soup.stripped_strings

def docx_add_text(doc, text, size, bold, alignment, font, first_line_indent=0.75, line_spacing=Pt(30)):
    text.strip().encode('utf-8')
    if text.__contains__('text-align: center'):
        alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for tmp in convert_to_plaintext_array(text):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
        p.paragraph_format.line_spacing = line_spacing
        run = p.add_run(tmp)
        run.font.size = Pt(size) # 12：小四号 16：三号
        run.bold = bold
        if alignment is not None:
            p.alignment = alignment
        run.font.name = font
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font)

def docx_add_title(doc, title):
    docx_add_text(doc, title, 16, True, WD_PARAGRAPH_ALIGNMENT.CENTER, '宋体')

def docx_add_author(doc, author):
    docx_add_text(doc, author, 12, False, WD_PARAGRAPH_ALIGNMENT.CENTER, '宋体')

def docx_add_subtitle(doc, subtitle):
    docx_add_text(doc, subtitle, 12, True, WD_PARAGRAPH_ALIGNMENT.CENTER, '宋体')

def docx_add_body(doc, body):
    docx_add_text(doc, body, 12, False, None, '仿宋')

def docx_add_image(doc, url, xy=0, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER):
    p = doc.add_paragraph()
    run = p.add_run()
    if xy > 1.5:
        run.add_picture(download_image(url), height=Inches(9))
    else:
        run.add_picture(download_image(url), width=Inches(6.5))
    if alignment is not None:
        p.alignment = alignment

def docx_add_link(doc, link):
    docx_add_text(doc, f'(链接：{link})', 10, False, WD_PARAGRAPH_ALIGNMENT.CENTER, '仿宋')

def convert_to_docx(article, path):
    title = get_title(article)
    author = get_author_name(article)
    content = get_content(article)
    create_time = get_create_time(article)

    if not os.path.exists(path):
        os.mkdir(path)
    
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    docx_add_title(doc, title)
    docx_add_author(doc, f'{author} {create_time.month:02d}-{create_time.day:02d}')

    # content
    for x in content:
        if x['type'] == 0:
            print(x)
            pass
        elif x['type'] == 1: # text/image
            has_img = x.keys().__contains__('img_url')
            has_text = x.keys().__contains__('text')
            has_height = x.keys().__contains__('img_height')
            has_width = x.keys().__contains__('img_width')
            if has_img:
                if has_text: # 1,1
                    docx_add_text(doc, x['text'], 12, False, None, '仿宋')
                else: # 1,0
                    pass
                if has_height and has_width:
                    xy = float(x['img_height']/x['img_width'])
                else:
                    xy = 0
                docx_add_image(doc, x['img_url'], xy)
            elif has_text: # 0,1
                docx_add_body(doc, x['text'])
            else: # 0,0
                logger.error(x)
        elif x['type'] == 2:
            print(x)
            pass
        elif x['type'] == 3: # video
            docx_add_image(doc, x['video_thumbnail'])
            docx_add_link(doc, x['video_url'])
        elif x['type'] == 4:
            print(x)
            pass
        elif x['type'] == 5:
            print(x)
            pass
        elif x['type'] == 6: # subtitle
            docx_add_subtitle(doc, x['text'])
        elif x['type'] == 7: # link
            docx_add_body(doc, x['text'])
        else:
            print(x)
            pass

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    doc.core_properties.author = "Administrator"
    doc.core_properties.comments = ""
    doc.save(f'{path}/{title}.docx')
    logger.success(f'文章保存成功   {path}/{title}.docx')

def load_task():
    tasks = []
    content = load_from_file(TASK_FILE)
    for line in content.split('\n'):
        line = line.removesuffix('\r')
        if not line.startswith('http'):
            logger.warning(f'无效的链接：{line}')
            continue
        if tasks.__contains__(line):
            logger.warning(f'重复的链接：{line}')
            continue
        tasks.append(line)
    logger.info(f'已读取{tasks.__len__()}条链接')
    return tasks

def allocate_tasks(tasks):
    total_cnt = tasks.__len__()
    if total_cnt < THREADS:
        global THREADS
        THREADS = total_cnt
        logger.warning(f'任务数({total_cnt})小于设定的线程数({THREADS})，重置线程数为任务数')

    ret = []
    for i in range(THREADS):
        ret.append([])
    
    for task_idx in range(total_cnt):
        for c in range(THREADS):
            if task_idx % THREADS == c:
                ret[c].append(tasks[task_idx])
                break

    m = f'分配给{THREADS}个线程执行，每个线程的任务数量为：'
    for t in ret:
        m += "[" + str(len(t)) + "] "
    logger.debug(m)

    return ret

def meipian_to_docx(urls, path):
    for url in urls:
        html = get_html(url)
        article = get_article(html)
        convert_to_docx(article, path)

def wait_until_compelete():
    while 1:
        cnt = threading.active_count()
        if cnt == 1:
            logger.success('全线程任务完成！')
            input('按任意键退出')
            break

def main():
    tasks = load_task()
    if tasks.__len__() == 0:
        logger.warning('无链接，程序终止')
        return
    task_lists = allocate_tasks(tasks)
    for i in range(len(task_lists)):
        thd = Thread(target=meipian_to_docx, args=(task_lists[i],OUTPUT_PATH,), name=f'Thread_{i}')
        thd.daemon = True
        thd.start()
    wait_until_compelete()

if __name__ == '__main__':
    # url = 'https://www.meipian.cn/4rzrnq2b'
    main()